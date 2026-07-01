"""Export solved papers to PDF and DOCX."""

from __future__ import annotations

import io
from typing import Any


def export_to_pdf(paper: dict[str, Any]) -> bytes:
    """Generate a printable PDF of a solved question paper."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    story: list[Any] = []

    # Custom styles
    title_style = ParagraphStyle(
        "PaperTitle", parent=styles["Title"], fontSize=16, spaceAfter=12, textColor=colors.HexColor("#0e7490")
    )
    section_style = ParagraphStyle(
        "SectionHeader", parent=styles["Heading2"], fontSize=13, spaceAfter=8, spaceBefore=16,
        textColor=colors.HexColor("#0e7490"), borderWidth=1, borderColor=colors.HexColor("#0e7490"),
        borderPadding=4,
    )
    question_style = ParagraphStyle(
        "Question", parent=styles["Normal"], fontSize=10, spaceAfter=4, leading=14,
        textColor=colors.HexColor("#1e293b"),
    )
    answer_style = ParagraphStyle(
        "Answer", parent=styles["Normal"], fontSize=10, spaceAfter=4, leading=14,
        textColor=colors.HexColor("#334155"), leftIndent=12,
    )
    key_point_style = ParagraphStyle(
        "KeyPoint", parent=styles["Normal"], fontSize=9, spaceAfter=2, leading=12,
        textColor=colors.HexColor("#047857"), leftIndent=24, bulletIndent=12,
    )
    tip_style = ParagraphStyle(
        "Tip", parent=styles["Normal"], fontSize=9, spaceAfter=6, leading=12,
        textColor=colors.HexColor("#b45309"), leftIndent=12, fontStyle="italic",
    )

    # ── Title ──
    paper_info = paper.get("paper_info", {})
    story.append(Paragraph("Question Paper — Solved", title_style))

    # Paper info table
    info_data = [
        ["Total Marks", str(paper_info.get("total_marks", "N/A"))],
        ["Duration", paper_info.get("duration", "N/A")],
        ["Source", paper.get("source", "N/A")],
    ]
    info_table = Table(info_data, colWidths=[4 * cm, 10 * cm])
    info_table.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#64748b")),
        ("TEXTCOLOR", (1, 0), (1, -1), colors.HexColor("#1e293b")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 12))

    # ── Questions & Answers ──
    solved = paper.get("solved_questions", [])
    current_section = None

    for q in solved:
        section = q.get("section", "All")
        if section != current_section:
            current_section = section
            story.append(Paragraph(f"Section {section}", section_style))

        marks = q.get("marks", 1)
        q_num = q.get("question_number", "")
        q_text = q.get("question_text", "")

        # Question line
        q_line = f"<b>Q{q_num}</b> <font color='#0e7490'>[{marks} mark{'s' if marks != 1 else ''}]</font> {q_text}"
        story.append(Paragraph(q_line, question_style))

        # Answer
        answer = q.get("answer", {})
        direct = answer.get("direct_answer", "")
        if direct:
            story.append(Paragraph(f"<b>Answer:</b> {direct}", answer_style))

        # Key points
        key_points = answer.get("key_points", [])
        if key_points:
            for kp in key_points:
                story.append(Paragraph(f"• {kp}", key_point_style))

        # Common mistakes
        mistakes = answer.get("common_mistakes", [])
        if mistakes:
            for m in mistakes:
                story.append(Paragraph(f"⚠ {m}", tip_style))

        # Exam tip
        tip = answer.get("exam_tips", "")
        if tip:
            story.append(Paragraph(f"Tip: {tip}", tip_style))

        story.append(Spacer(1, 6))

    # ── Pattern Analysis (if present) ──
    pattern = paper.get("pattern_analysis", {})
    if pattern:
        story.append(PageBreak())
        story.append(Paragraph("Pattern Analysis", title_style))

        diff = pattern.get("difficulty_distribution", {})
        if diff:
            diff_data = [["Difficulty", "Percentage"]] + [[k, f"{v}%"] for k, v in diff.items()]
            diff_table = Table(diff_data, colWidths=[5 * cm, 5 * cm])
            diff_table.setStyle(TableStyle([
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0e7490")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(Paragraph("<b>Difficulty Distribution</b>", question_style))
            story.append(diff_table)
            story.append(Spacer(1, 12))

        topics = pattern.get("topic_weightage", {})
        if topics:
            topic_data = [["Topic", "Weightage"]] + [[k, f"{v}%"] for k, v in list(topics.items())[:10]]
            topic_table = Table(topic_data, colWidths=[8 * cm, 4 * cm])
            topic_table.setStyle(TableStyle([
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0e7490")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(Paragraph("<b>Topic Weightage</b>", question_style))
            story.append(topic_table)

    doc.build(story)
    return buffer.getvalue()


def export_to_docx(paper: dict[str, Any]) -> bytes:
    """Generate an editable DOCX of a solved question paper."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("Question Paper — Solved", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Paper info
    paper_info = paper.get("paper_info", {})
    info_para = doc.add_paragraph()
    info_para.add_run("Total Marks: ").bold = True
    info_para.add_run(f"{paper_info.get('total_marks', 'N/A')}    ")
    info_para.add_run("Duration: ").bold = True
    info_para.add_run(f"{paper_info.get('duration', 'N/A')}    ")
    info_para.add_run("Source: ").bold = True
    info_para.add_run(paper.get("source", "N/A"))

    doc.add_paragraph("")  # spacer

    # ── Questions & Answers ──
    solved = paper.get("solved_questions", [])
    current_section = None

    for q in solved:
        section = q.get("section", "All")
        if section != current_section:
            current_section = section
            sec_heading = doc.add_heading(f"Section {section}", level=2)
            for run in sec_heading.runs:
                run.font.color.rgb = RGBColor(0x0e, 0x74, 0x90)

        marks = q.get("marks", 1)
        q_num = q.get("question_number", "")
        q_text = q.get("question_text", "")

        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{q_num} [{marks} mark{'s' if marks != 1 else ''}] ")
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(0x0e, 0x74, 0x90)
        q_para.add_run(q_text)

        # Answer
        answer = q.get("answer", {})
        direct = answer.get("direct_answer", "")
        if direct:
            ans_para = doc.add_paragraph()
            ans_para.add_run("Answer: ").bold = True
            ans_para.add_run(direct)
            ans_para.paragraph_format.left_indent = Inches(0.3)

        # Key points
        key_points = answer.get("key_points", [])
        if key_points:
            kp_heading = doc.add_paragraph()
            kp_heading.add_run("Key Points:").bold = True
            kp_heading.paragraph_format.left_indent = Inches(0.3)
            for kp in key_points:
                bp = doc.add_paragraph(kp, style="List Bullet")
                bp.paragraph_format.left_indent = Inches(0.5)

        # Common mistakes
        mistakes = answer.get("common_mistakes", [])
        if mistakes:
            m_heading = doc.add_paragraph()
            m_heading.add_run("Common Mistakes:").bold = True
            m_heading.paragraph_format.left_indent = Inches(0.3)
            for m in mistakes:
                bp = doc.add_paragraph(m, style="List Bullet")
                bp.paragraph_format.left_indent = Inches(0.5)

        # Exam tip
        tip = answer.get("exam_tips", "")
        if tip:
            tip_para = doc.add_paragraph()
            tip_para.add_run("Tip: ").bold = True
            tip_para.add_run(tip)
            tip_para.paragraph_format.left_indent = Inches(0.3)
            for run in tip_para.runs:
                run.font.color.rgb = RGBColor(0xb4, 0x53, 0x09)

        doc.add_paragraph("")  # spacer

    # ── Pattern Analysis ──
    pattern = paper.get("pattern_analysis", {})
    if pattern:
        doc.add_page_break()
        pa_title = doc.add_heading("Pattern Analysis", level=1)
        for run in pa_title.runs:
            run.font.color.rgb = RGBColor(0x0e, 0x74, 0x90)

        diff = pattern.get("difficulty_distribution", {})
        if diff:
            doc.add_heading("Difficulty Distribution", level=2)
            for k, v in diff.items():
                doc.add_paragraph(f"{k}: {v}%", style="List Bullet")

        topics = pattern.get("topic_weightage", {})
        if topics:
            doc.add_heading("Topic Weightage", level=2)
            for k, v in list(topics.items())[:10]:
                doc.add_paragraph(f"{k}: {v}%", style="List Bullet")

        high_value = pattern.get("high_value_topics", [])
        if high_value:
            doc.add_heading("High Value Topics", level=2)
            for t in high_value:
                doc.add_paragraph(t, style="List Bullet")

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_to_text(paper: dict[str, Any]) -> str:
    """Export solved paper as plain text."""
    lines: list[str] = []

    paper_info = paper.get("paper_info", {})
    lines.append("QUESTION PAPER — SOLVED")
    lines.append("=" * 50)
    lines.append(f"Total Marks: {paper_info.get('total_marks', 'N/A')}")
    lines.append(f"Duration: {paper_info.get('duration', 'N/A')}")
    lines.append(f"Source: {paper.get('source', 'N/A')}")
    lines.append("")

    solved = paper.get("solved_questions", [])
    current_section = None

    for q in solved:
        section = q.get("section", "All")
        if section != current_section:
            current_section = section
            lines.append(f"\nSECTION {section}")
            lines.append("-" * 30)

        marks = q.get("marks", 1)
        q_num = q.get("question_number", "")
        q_text = q.get("question_text", "")

        lines.append(f"Q{q_num} [{marks} mark{'s' if marks != 1 else ''}] {q_text}")

        answer = q.get("answer", {})
        direct = answer.get("direct_answer", "")
        if direct:
            lines.append(f"  Answer: {direct}")

        key_points = answer.get("key_points", [])
        if key_points:
            lines.append("  Key Points:")
            for kp in key_points:
                lines.append(f"    - {kp}")

        tip = answer.get("exam_tips", "")
        if tip:
            lines.append(f"  Tip: {tip}")

        lines.append("")

    return "\n".join(lines)
